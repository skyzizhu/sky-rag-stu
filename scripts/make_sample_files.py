#!/usr/bin/env python3
"""生成 6 种格式的示例知识文件到 knowledge/ 目录。

项目自带了一批示例文件（示例_开头的 6 个），方便第一次就能跑通全链路。
这个脚本用来（重新）生成其中 3 个「二进制类」文件：docx / pdf / rtf。
其余 3 个文本类示例（md / txt / html）直接以文件形式放在 knowledge/ 里。

想删掉所有示例重新开始：直接删除 knowledge/ 下的「示例_」文件即可。
"""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
KNOWLEDGE = ROOT / "knowledge"


# ---------------------------------------------------------------- RTF（纯文本生成，中文用 \uN 转义）
def _rtf_escape(text: str) -> str:
    out = []
    for ch in text:
        if ord(ch) > 127:
            out.append(f"\\u{ord(ch)}?")
        elif ch in "\\{}":
            out.append("\\" + ch)
        else:
            out.append(ch)
    return "".join(out)


def make_rtf() -> None:
    paragraphs = [
        ("产品研发周会纪要",),
        ("时间：2026年8月20日 10:00-11:00", "参会人：产品组全员、研发代表"),
        ("一、本周进展", "1. 知识库搜索功能灰度上线，覆盖 20% 用户。", "2. 文档中心改版完成视觉验收。"),
        ("二、问题与风险", "1. 搜索无结果率仍高达 18%，主要因为用户提问过于口语化。", "2. 移动端首屏加载时间超过 3 秒，需性能专项优化。"),
        ("三、下周计划", "1. 上线同义词扩展，目标把无结果率降到 10% 以内。", "2. 启动性能专项，目标首屏降到 1.5 秒。"),
        ("四、决议", "自下月起周会纪要统一归档到知识库，命名格式：产品研发周会纪要_YYYYMMDD。"),
    ]
    body = r"{\rtf1\ansi\deff0{\fonttbl{\f0 PingFang SC;}}"
    for para in paragraphs:
        for line in para:
            body += r"\par " + _rtf_escape(line)
        body += r"\par"
    body += "}"

    path = KNOWLEDGE / "reference" / "documents" / "示例_会议纪要.rtf"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(body, encoding="utf-8")
    print(f"✅ {path.relative_to(ROOT).as_posix()}")


# ---------------------------------------------------------------- DOCX（python-docx）
def make_docx() -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("云笔记 App 产品需求文档 V1.0", level=0)

    doc.add_heading("一、背景与目标", level=1)
    doc.add_paragraph(
        "用户反馈碎片化信息难以沉淀，需要在云笔记 App 中提供轻量的知识管理能力。"
        "本期目标是上线「标签 + 全文搜索」组合，让用户 3 步内找到 3 个月内的任意一条笔记。"
    )

    doc.add_heading("二、功能需求", level=1)
    doc.add_paragraph("1. 支持给笔记打标签，单条笔记最多 10 个标签。")
    doc.add_paragraph("2. 全文搜索支持标题和正文，结果按修改时间排序。")
    doc.add_paragraph("3. 回收站内容保留 30 天，到期自动清空。")

    doc.add_heading("三、验收标准", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    head = ["指标", "目标值", "说明"]
    rows = [
        ["搜索成功率", "≥ 90%", "用户搜索后 3 次点击内打开目标笔记"],
        ["搜索响应时间", "≤ 800ms", "P95，1 万条笔记以内"],
        ["标签使用率", "≥ 40%", "周活跃用户中使用过标签的比例"],
    ]
    for col, text in enumerate(head):
        table.rows[0].cells[col].text = text
    for row_index, row in enumerate(rows, start=1):
        for col, text in enumerate(row):
            table.rows[row_index].cells[col].text = text

    path = KNOWLEDGE / "reference" / "documents" / "示例_产品需求.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"✅ {path.relative_to(ROOT).as_posix()}")


# ---------------------------------------------------------------- PDF（手写极简 PDF 生成器，演示页眉/页脚清理和页码追溯）
def _pdf_escape(text: str) -> str:
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def _build_simple_pdf(pages: list[list[str]]) -> bytes:
    """生成只含英文文本的多页 PDF（内置 Helvetica 字体不支持中文）。"""
    objects: list[bytes] = []

    def add(obj: bytes) -> int:
        objects.append(obj)
        return len(objects)

    def content_stream(lines: list[str]) -> bytes:
        parts = ["BT", "/F1 12 Tf", "72 720 Td", "16 TL"]
        for line in lines:
            parts.append(f"({_pdf_escape(line)}) Tj T*")
        parts.append("ET")
        return "\n".join(parts).encode("latin-1")

    # 对象编号：1=Catalog, 2=Pages, 3=Font；每页占 2 个对象（Page + Contents）
    page_ids = []
    content_ids = []
    for _ in pages:
        content_ids.append(add(b""))          # 占位，稍后回填
        page_ids.append(add(b""))             # 占位

    catalog_id = add(b"")                     # Catalog
    pages_id = add(b"")                       # Pages
    font_id = add(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    for i, lines in enumerate(pages):
        stream = content_stream(lines)
        objects[content_ids[i] - 1] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream"
        )
        objects[page_ids[i] - 1] = (
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 612 792] "
            f"/Contents {content_ids[i]} 0 R /Resources << /Font << /F1 {font_id} 0 R >> >> >>"
        ).encode()

    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    objects[catalog_id - 1] = f"<< /Type /Catalog /Pages {pages_id} 0 R >>".encode()
    objects[pages_id - 1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode()

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0] * (len(objects) + 1)
    for index, obj in enumerate(objects, start=1):
        offsets[index] = len(out)
        out += f"{index} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_pos = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_id} 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF"
    ).encode()
    return bytes(out)


def make_pdf() -> None:
    header = "SkyNotes Product Spec - Internal Use Only"  # 每页重复的页眉，用来演示清洗节点
    pages = [
        [header, "", "SkyNotes Release Notes - Page 1", "",
         "Feature A: Markdown export was shipped in March 2026.",
         "Feature A supports tables, code blocks and footnotes.",
         "",
         "Feature B: Tag management entered beta in May 2026.",
         "Each note can carry up to 20 tags in the beta build."],
        [header, "", "SkyNotes Release Notes - Page 2", "",
         "Feature C: Offline mode is planned for Q4 2026.",
         "Offline mode will cache the most recent 500 notes per device.",
         "",
         "Team note: all release dates in this document follow the",
         "roadmap approved on 2026-01-15 and may be adjusted."],
        [header, "", "SkyNotes Release Notes - Page 3", "",
         "Feature D: AI summaries rolled out to all users in July 2026.",
         "Summaries are limited to 3 bullet points per note.",
         "",
         "Contact: knowledge-team@skynotes.example.com"],
    ]
    path = KNOWLEDGE / "reference" / "documents" / "示例_英文产品说明.pdf"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(_build_simple_pdf(pages))
    print(f"✅ {path.relative_to(ROOT).as_posix()}")


def make_rag_reference_pdf() -> None:
    """外部 RAG 参考资料（英文，多页，带重复页眉）。用于 Domain 过滤测试。"""
    header = "External Reading Collection - Personal Knowledge Base"
    pages = [
        [header, "", "External Reading: Retrieval-Augmented Generation - Page 1", "",
         "Retrieval-Augmented Generation (RAG) combines a retrieval module",
         "with a text generation model. The retrieval module searches an",
         "external knowledge base for passages relevant to the user query.",
         "",
         "Key benefit: answers are grounded in source documents, which",
         "reduces hallucination and enables citation of sources."],
        [header, "", "External Reading: Retrieval-Augmented Generation - Page 2", "",
         "A vector database stores text embeddings and supports fast",
         "similarity search. Popular choices include Qdrant, Milvus and",
         "Weaviate. Qdrant supports payload filtering, which allows the",
         "system to restrict search results by metadata fields such as",
         "domain, category, topic and status.",
         "",
         "Chunking strategy matters: chunks that are too small lose",
         "context, while chunks that are too large dilute relevance."],
        [header, "", "External Reading: Retrieval-Augmented Generation - Page 3", "",
         "Recommended default parameters for a personal knowledge base:",
         "chunk size 500 to 800 characters, overlap 80 to 150 characters,",
         "and top-k between 3 and 5. Always evaluate retrieval quality",
         "with a question set before tuning any further.",
         "",
         "Source: external articles collected in August 2026."],
    ]
    path = KNOWLEDGE / "reference" / "articles" / "RAG外部资料.pdf"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(_build_simple_pdf(pages))
    print(f"✅ {path.relative_to(ROOT).as_posix()}")


if __name__ == "__main__":
    KNOWLEDGE.mkdir(exist_ok=True)
    make_rtf()
    make_docx()
    make_pdf()
    make_rag_reference_pdf()
    print("\n文本类示例（md / txt / html）已随项目提供，无需生成。")
    sys.exit(0)

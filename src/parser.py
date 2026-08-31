"""节点 1：文档解析（Parser）。

产品视角：这一步相当于「收发室」——
不管你递进来的是 Word、PDF、网页剪藏还是纯文本，
它都负责把内容完整读出来，并登记一张「档案卡」（元数据 Metadata），
后续所有节点都只认这种统一的格式。

统一输出结构：
    {
        "text": "文档正文",
        "metadata": {
            "source":      文件名（来源）
            "file_type":   文件类型
            "title":       标题
            "section":     章节（解析阶段为空，切片时才知道）
            "page":        页码（解析阶段为空，切片时才知道）
            "category":    分类
            "created_at":  文件创建日期
            "updated_at":  文件修改日期
            "version":     版本（V1 固定 1.0）
            "status":      状态（V1 固定 active）
        }
    }
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader
from striprtf.striprtf import rtf_to_text

from src.config import get_config
from src.metadata import parse_front_matter

# 支持的文件格式
SUPPORTED_EXTENSIONS = {".txt", ".md", ".rtf", ".html", ".htm", ".docx", ".pdf"}

# PDF 分页标记：解析 PDF 时插入正文，用来记住「这段话原来在第几页」。
# 切片节点会读取这个标记，把页码写进元数据。
PAGE_MARKER_TEMPLATE = "<<<__RAG_PAGE__={page}>>>"
PAGE_MARKER_RE = re.compile(r"^<<<__RAG_PAGE__=(\d+)>>>\s*$", re.MULTILINE)


class ParseError(Exception):
    """某个文件解析失败时抛出，信息里带文件名。"""


@dataclass
class ParsedDocument:
    """内部统一的文档结构：正文 + 元数据 + Front Matter。"""

    text: str
    metadata: dict = field(default_factory=dict)
    front_matter: dict = field(default_factory=dict)  # Markdown 头部声明（非 md 文件为空）


def _file_dates(path: Path) -> tuple[str, str]:
    """读取文件的创建日期和修改日期（只保留到「日」）。"""
    stat = path.stat()
    created = getattr(stat, "st_birthtime", None) or stat.st_ctime
    updated = stat.st_mtime
    created_at = datetime.fromtimestamp(created).strftime("%Y-%m-%d")
    updated_at = datetime.fromtimestamp(updated).strftime("%Y-%m-%d")
    return created_at, updated_at


def _base_metadata(path: Path, title: str, file_type: str) -> dict:
    cfg = get_config()
    created_at, updated_at = _file_dates(path)
    return {
        "source": path.name,
        "file_type": file_type,
        "title": title,
        "section": None,
        "page": None,
        "category": cfg.default_category,
        "created_at": created_at,
        "updated_at": updated_at,
        "version": "1.0",
        "status": "active",
    }


def _read_txt(path: Path) -> tuple[str, str, dict]:
    """纯文本：直接读，标题用文件名。"""
    return path.read_text(encoding="utf-8", errors="replace"), "", {}


def _read_markdown(path: Path) -> tuple[str, str, dict]:
    """Markdown：先拆 Front Matter（人工声明的分类信息），标题取正文的第一个一级标题。"""
    raw = path.read_text(encoding="utf-8", errors="replace")
    front_matter, text = parse_front_matter(raw)
    match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
    title = match.group(1).strip() if match else ""
    return text, title, front_matter


def _read_rtf(path: Path) -> tuple[str, str, dict]:
    """RTF：用专门的解析库读出纯文本。"""
    raw = path.read_text(encoding="utf-8", errors="replace")
    return rtf_to_text(raw), "", {}


def _read_html(path: Path) -> tuple[str, str, dict]:
    """HTML：提取正文区域，并尽量保留标题、列表、表格、引用与代码结构。"""
    # 直接交给 BeautifulSoup 读取字节，可依据 BOM / meta charset 识别旧中文网页编码。
    soup = BeautifulSoup(path.read_bytes(), "html.parser")
    for tag in soup([
        "script", "style", "noscript", "template", "svg", "canvas", "iframe",
        "button", "input", "select", "textarea",
    ]):
        tag.decompose()

    h1 = soup.find("h1")
    page_title = soup.title.get_text(" ", strip=True) if soup.title else ""
    title = h1.get_text(" ", strip=True) if h1 else page_title

    body = soup.body or soup
    candidates = body.select("article, main, [role='main']")
    root = max(candidates, key=lambda node: len(node.get_text(" ", strip=True)), default=body)
    body_length = len(body.get_text(" ", strip=True))
    root_length = len(root.get_text(" ", strip=True))
    # 很短的 main/article 可能只是推荐卡片，不应覆盖真正的 body 正文。
    if root is not body and root_length < 120 and root_length < body_length * 0.35:
        root = body

    noise_tags = ["nav", "footer", "aside", "form", "dialog"]
    if root is body:
        noise_tags.append("header")
    for tag in root.find_all(noise_tags):
        tag.decompose()
    for tag in root.select("[hidden], [aria-hidden='true']"):
        tag.decompose()

    # 先转换会包含其他标签的结构，再处理普通块级元素。
    for table in list(root.find_all("table")):
        rows = []
        for row in table.find_all("tr"):
            cells = [
                re.sub(r"\s+", " ", cell.get_text(" ", strip=True)).replace("|", "\\|")
                for cell in row.find_all(["th", "td"])
            ]
            if cells:
                rows.append(cells)
        if rows:
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            table_lines = ["| " + " | ".join(row) + " |" for row in rows]
            table_lines.insert(1, "| " + " | ".join(["---"] * width) + " |")
            table.replace_with("\n\n" + "\n".join(table_lines) + "\n\n")
        else:
            table.decompose()

    for pre in list(root.find_all("pre")):
        code = pre.get_text("\n", strip=False).strip("\n")
        pre.replace_with(f"\n\n```\n{code}\n```\n\n")
    for code in list(root.find_all("code")):
        code.replace_with(f"`{code.get_text('', strip=True)}`")

    for image in list(root.find_all("img")):
        alt = re.sub(r"\s+", " ", image.get("alt", "")).strip()
        image.replace_with(f"[图片：{alt}]" if alt else "")
    for anchor in list(root.find_all("a")):
        label = re.sub(r"\s+", " ", anchor.get_text(" ", strip=True))
        href = (anchor.get("href") or "").strip()
        if label and href.startswith(("http://", "https://")):
            anchor.replace_with(f"[{label}]({href})")
        else:
            anchor.replace_with(label)

    for list_tag in reversed(root.find_all(["ol", "ul"])):
        lines = []
        for index, item in enumerate(list_tag.find_all("li", recursive=False), start=1):
            item_text = re.sub(r"\s+", " ", item.get_text(" ", strip=True))
            if item_text:
                marker = f"{index}." if list_tag.name == "ol" else "-"
                lines.append(f"{marker} {item_text}")
        list_tag.replace_with("\n\n" + "\n".join(lines) + "\n\n")

    for quote in list(root.find_all("blockquote")):
        lines = [line.strip() for line in quote.get_text("\n").splitlines() if line.strip()]
        quote.replace_with("\n\n" + "\n".join(f"> {line}" for line in lines) + "\n\n")
    for heading in list(root.find_all(re.compile(r"^h[1-6]$"))):
        level = int(heading.name[1])
        heading.replace_with(f"\n\n{'#' * level} {heading.get_text(' ', strip=True)}\n\n")
    for br in list(root.find_all("br")):
        br.replace_with("\n")
    for block in root.find_all(["p", "div", "section", "article", "main", "header", "address", "figure", "figcaption", "dl", "dt", "dd"]):
        block.insert_before("\n\n")
        block.insert_after("\n\n")

    text = root.get_text("", strip=False)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text, title, {}


def _read_docx(path: Path) -> tuple[str, str, dict]:
    """Word：按段落读正文；标题样式段落转成 Markdown 风格的 # 标记，
    让切片节点能识别章节（Heading 1 → #，Heading 2 → ## ……）。表格逐行读。"""
    doc = DocxDocument(str(path))
    parts: list[str] = []
    title = ""
    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            parts.append("")
            continue
        style = para.style.name if para.style is not None else ""
        if style == "Title" and not title:
            title = line
        elif style.startswith("Heading"):
            try:
                level = min(int(style.split()[-1]), 6)
            except ValueError:
                level = 2
            parts.append("#" * level + " " + line)
            if not title:
                title = line
        else:
            parts.append(line)
    for table in doc.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        parts.append("\n".join(rows))
    return "\n".join(parts), title, {}


def _read_pdf(path: Path) -> tuple[str, str, dict]:
    """PDF：逐页读，页与页之间插入分页标记（页码信息全靠它往后传）。"""
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ParseError("PDF 已加密，无法读取")
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        pages.append(f"{PAGE_MARKER_TEMPLATE.format(page=index)}\n{page_text}")
    return "\n".join(pages), "", {}


_READERS = {
    ".txt": _read_txt,
    ".md": _read_markdown,
    ".rtf": _read_rtf,
    ".html": _read_html,
    ".htm": _read_html,
    ".docx": _read_docx,
    ".pdf": _read_pdf,
}


def parse_file(path: Path | str) -> ParsedDocument:
    """把单个文件解析成统一的 ParsedDocument。失败会抛出带文件名的错误。"""
    path = Path(path)
    if not path.exists():
        raise ParseError(f"文件不存在：{path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        supported = "、".join(sorted(SUPPORTED_EXTENSIONS))
        raise ParseError(f"暂不支持 {ext} 格式（目前支持：{supported}）：{path.name}")

    try:
        text, title, front_matter = _READERS[ext](path)
    except ParseError:
        raise
    except Exception as exc:  # 解析库的各类异常统一包装，避免用户看到天书
        raise ParseError(f"解析失败：{path.name}（{type(exc).__name__}: {exc}）") from exc

    metadata = _base_metadata(path, title or path.stem, ext.lstrip("."))
    return ParsedDocument(text=text, metadata=metadata, front_matter=front_matter)


def parse_directory(dir_path: Path | str) -> tuple[list[ParsedDocument], list[str]]:
    """解析整个目录（含子目录），返回（成功列表, 失败原因列表）。"""
    dir_path = Path(dir_path)
    if not dir_path.exists():
        raise ParseError(f"知识目录不存在：{dir_path}")

    docs: list[ParsedDocument] = []
    failures: list[str] = []
    for path in sorted(dir_path.rglob("*")):
        if not path.is_file() or path.name.startswith("."):
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            docs.append(parse_file(path))
        except ParseError as exc:
            failures.append(str(exc))
    return docs, failures


# ---------------- 直接运行本文件：打印解析结果，人工检查 ----------------
if __name__ == "__main__":
    cfg = get_config()
    docs, failures = parse_directory(cfg.knowledge_dir)
    print(f"知识目录：{cfg.knowledge_dir}")
    print(f"解析成功 {len(docs)} 个文件，失败 {len(failures)} 个\n")

    for failure in failures:
        print(f"❌ {failure}")

    for doc in docs:
        meta = doc.metadata
        preview = re.sub(r"\s+", " ", doc.text[:120])
        print("=" * 60)
        print(f"文件：{meta['source']}（{meta['file_type']}）")
        print(f"标题：{meta['title']}  字数：{len(doc.text)}")
        print(f"日期：创建 {meta['created_at']} / 修改 {meta['updated_at']}")
        print(f"正文开头：{preview}……")
    sys.exit(0)
